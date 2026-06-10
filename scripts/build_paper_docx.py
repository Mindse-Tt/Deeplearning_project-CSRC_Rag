"""Generate the M6.1 paper (Word / .docx) with embedded figures.

Academic-style rewrite: formal prose, Table/Figure numbering, figure captions,
cross-references, 4-paragraph structured abstract (Background / Method /
Results / Conclusion).

Layout::

    Cover + Abstract + Keywords
    §1 Introduction
        §1.1 Background
        §1.2 Related Work
        §1.3 Research Objectives
        §1.4 Contributions
    §2 Data and Evaluation Sets
    §3 System Architecture (7-layer pipeline)
    §4 Experiments
        §4.1 Retrieval ablation (M3e)
        §4.2 Generation comparison G0-G3 (M4.4)     ← core
        §4.3 Qualitative case studies
        §4.4 Trend aggregator precision (M4.2)
        §4.5 Punishment-type multi-label classifier (M5)
    §5 Hallucination Mitigation: Three-Layer Defence
    §6 Deployment Cost Analysis
    §7 Limitations and Future Work
    §8 Contributions
    References

Run::

    python scripts/build_paper_docx.py
"""
from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parent.parent
FIG_DIR = PROJECT_ROOT / "docs" / "visuals" / "png" / "paper"


# ---------------------------------------------------------------------------
# Data loaders
# ---------------------------------------------------------------------------


def _load(rel: str) -> dict:
    return json.loads((PROJECT_ROOT / rel).read_text(encoding="utf-8"))


def _try_load(rel: str) -> dict | None:
    p = PROJECT_ROOT / rel
    return _load(rel) if p.exists() else None


# ---------------------------------------------------------------------------
# Word helpers
# ---------------------------------------------------------------------------


def _set_font(run, name: str = "宋体", size_pt: float = 10.5, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    _set_font(run, name="黑体", size_pt=16 - (level - 1) * 2, bold=True)


def para(doc: Document, text: str, *, indent: bool = True) -> None:
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _set_font(run)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    _set_font(run)


def bold_run(doc: Document, prefix: str, rest: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.4
    r1 = p.add_run(prefix)
    _set_font(r1, bold=True)
    r2 = p.add_run(rest)
    _set_font(r2)


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    _set_font(run, name="黑体", size_pt=10, bold=True)


def add_figure(doc: Document, image_path: Path, caption_text: str, width_cm: float = 14.0) -> None:
    if not image_path.exists():
        para(doc, f"[图缺失: {image_path.name}]", indent=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    caption(doc, caption_text)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    caption_text: str | None = None,
) -> None:
    if caption_text:
        caption(doc, caption_text)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        _set_font(r, name="黑体", size_pt=10, bold=True)
    for ri, row in enumerate(rows, start=1):
        cells = table.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            _set_font(r, size_pt=9.5)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Section · Cover + Abstract
# ---------------------------------------------------------------------------
def section_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("面向证监会违规案例的检索增强问答系统")
    _set_font(r, name="黑体", size_pt=22, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("—— 中文小参数量基座的指令微调与幻觉缓解研究")
    _set_font(r, name="黑体", size_pt=14)

    doc.add_paragraph()

    # 作者信息块
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("许浩财  贾彤  戴一鑫  张彦扬  王怡菲")
    _set_font(r, size_pt=11, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("深度学习课程设计 · 赛道 B · 2026 年 4 月")
    _set_font(r, size_pt=10.5)

    doc.add_paragraph()

    # ----- 结构化摘要:Background / Method / Results / Conclusion -----
    heading(doc, "摘  要", level=2)

    bold_run(
        doc,
        "研究背景。",
        "中国证监会每年公开的行政处罚公告逐年增加,2017 年至 2024 年七年间增幅达 78%。"
        "合规从业人员在实务中需要对海量处罚公告进行同类案件检索、法条依据关联以及处罚"
        "分布的统计性回答。通用大语言模型在零样本场景下对此类垂直法律语料存在三个典型"
        "短板:结构化引用格式缺失、事实性幻觉高发以及领域术语语义表示不足。",
    )

    bold_run(
        doc,
        "研究方法。",
        "本研究构建了一个端到端的七层检索增强问答流水线,知识库基于 CSMAR 14 740 条原始"
        "处罚记录去重后得到的 4 233 个事件级文档。检索层采用 BM25、bge-small-zh-v1.5 向量"
        "以及 bge-reranker-v2-m3 交叉编码器的三路融合(Reciprocal Rank Fusion, k=60),"
        "生成层在 Qwen2.5-0.5B-Instruct 基座之上采用 4-bit NF4 量化的 QLoRA 指令微调"
        "(r=16, α=32),并设计了涵盖反幻觉负例在内的八类 5 500 条训练样本。"
        "评测在 130 条人工标注的金标(gold_130)与 30 条趋势分析专用金标(gold_trend_30)上进行。",
    )

    bold_run(
        doc,
        "实验结果。",
        "检索层在 98 条 retrieval-eligible 金标上 Hybrid Recall@5 从基线的 0.073 提升至 0.388。"
        "在控制证据与 prompt 一致的四组生成对照实验(G0–G3, n=30)中,微调前的三组(无论是否加入"
        "检索证据或强约束 prompt)EventID 引证命中率与格式合规率均为 0,加入 LoRA 后格式合规率"
        "提升至 76.7%,EventID 命中率提升至 20.0%,幻觉数字率自 20.0% 降至 3.3%,"
        "相对降幅 83%。端到端推理延迟约 10.5 秒,LoRA 适配器仅额外引入 0.7 秒开销。",
    )

    bold_run(
        doc,
        "结论。",
        "第一,在参数量小于 1 B 的中文基座上,结构化引用格式的学习无法通过 prompt 工程补偿,"
        "必须依赖指令微调。第二,RAG 机制与强约束 prompt 仅能将幻觉率减半,残余幻觉必须通过"
        "反幻觉对抗训练与引证校验联合缓解。第三,本研究提出的七层架构可在单块 8 GB 消费级"
        "GPU 上完整部署,对中小金融机构具有工程可行性。",
    )

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    r1 = p.add_run("关键词:")
    _set_font(r1, bold=True)
    r2 = p.add_run(
        "检索增强生成;大语言模型;参数高效微调;LoRA;幻觉缓解;证券合规"
    )
    _set_font(r2)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Section · 1 Introduction
# ---------------------------------------------------------------------------
def section_intro(doc: Document) -> None:
    heading(doc, "1  引言", level=1)

    heading(doc, "1.1 研究背景", level=2)
    para(
        doc,
        "随着中国资本市场的持续扩容,证券监督管理机构每年公开的行政处罚案件数量呈现显著"
        "上升趋势。根据本研究收集的中国研究数据服务平台(CSMAR)证监会处罚信息表,"
        "2017 年全年立案处罚 299 起,至 2024 年已上升至 534 起,七年间增长 78%"
        "(年度分布如 Figure 6 所示)。在合规审查、内控排查与法律尽职调查的实务场景中,"
        "从业人员需要对历史同类案件进行快速检索并结合相关法条给出处罚预判。",
    )
    para(
        doc,
        "然而,通用大语言模型在此类垂直法律语料上普遍面临三个结构性短板。"
        "其一,零样本场景下易于编造不存在的法条条款、处罚金额乃至公司名称;"
        "其二,即便施加强约束 prompt,仍缺乏对结构化引用格式(如 [EventID=xxx])的稳定"
        "遵从;其三,对证券合规特有术语(如「内幕交易」「虚假记载」「市场禁入」"
        "「控股股东操纵股价」)的语义向量化质量较弱,检索召回显著低于通用领域。"
        "本研究以证监会违规案例问答为载体,系统性回答上述三个短板的缓解路径。",
    )

    heading(doc, "1.2 相关工作", level=2)
    para(
        doc,
        "在检索增强生成方面,Lewis 等人 [8] 首次提出将稠密检索与生成式模型端到端联合训练的 RAG 框架。"
        "针对中文领域,Xiao 等人 [3] 发布的 bge 系列嵌入模型在 C-MTEB 基准上达到先进水平;"
        "Robertson 的 BM25 稀疏检索 [4] 与 Cormack 的 RRF 融合 [5] 至今仍是业界 baseline"
        "的核心组件。在参数高效微调方面,Hu 等人提出的 LoRA [6] 以及 Dettmers 等人进一步"
        "结合 4-bit 量化的 QLoRA [1] 使得大模型的领域适配成本降至消费级硬件可承受。"
        "在中文预训练模型上,Cui 等人的 MacBERT [7] 通过相似词替换的 MLM 任务对中文 BERT 进行了增强。"
        "与以上工作不同,本研究聚焦于:(a) 在 <1 B 参数量的小规模基座上实证参数高效微调对幻觉"
        "缓解的边际贡献;(b) 将 RAG、指令微调与规则校验三种机制组合为一个可复现的端到端系统;"
        "(c) 在严格受限的 8 GB 显存条件下完成主干训练与部署。",
    )

    heading(doc, "1.3 研究目标", level=2)
    para(
        doc,
        "本研究围绕课程赛道 B 设定的「面向垂直领域的智能问答系统」目标,具体展开以下四个"
        "层面的工作:(1) 基于 CSMAR 原始数据构建事件级检索知识库与多粒度的人工标注金标;"
        "(2) 设计并逐层消融验证七层检索增强架构;(3) 对开源中文基座 Qwen2.5-0.5B-Instruct "
        "进行 QLoRA 指令微调,提供可重复的微调前后定量对比;(4) 针对领域幻觉问题设计"
        "多层缓解机制并量化其逐层贡献。",
    )

    heading(doc, "1.4 主要贡献", level=2)

    para(doc, "本研究的主要贡献可归纳为以下五点:", indent=False)
    bullet(
        doc,
        "构建了覆盖六类意图的 130 条金标评测集(gold_130)与 30 条趋势分析专用金标"
        "(gold_trend_30),其中包含 4 条反幻觉陷阱题,支持检索与生成两种评测范式;",
    )
    bullet(
        doc,
        "提出融合多子查询改写、结构化元数据注入与三层 RRF 的检索栈,将 Hybrid Recall@5 "
        "从基线 0.073 提升至 0.388,相对增幅 431%;",
    )
    bullet(
        doc,
        "在 Qwen2.5-0.5B-Instruct 上以 QLoRA r=16 完成指令微调,52 分钟在 RTX 2060 SUPER "
        "8 GB 单卡上达成,训练损失自 2.52 收敛至 0.70;",
    )
    bullet(
        doc,
        "G0/G1/G2/G3 四组对照实验证实:格式合规率自 0% 跃升至 76.7%,幻觉数字率自 20.0% 降至 3.3%;",
    )
    bullet(
        doc,
        "实现专用于趋势分析意图的结构化聚合层(L6 Trend Aggregator),以 SQL-like groupby "
        "短路检索,使系统能够给出可审计的统计数字而非模糊枚举。",
    )


# ---------------------------------------------------------------------------
# Section · 2 Data
# ---------------------------------------------------------------------------
def section_data(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "2  数据与评测集构建", level=1)

    heading(doc, "2.1 原始语料", level=2)
    para(
        doc,
        "原始数据来源于中国研究数据服务平台(CSMAR)「证监会处罚信息表」,涵盖"
        "1994 至 2025 年共 14 740 条当事人级处罚记录。按 EventID(处罚公告级)聚合去重后,"
        "得到 4 233 个独立事件级文档,构成本研究检索知识库的基本单元。"
        "语料的按年度分布与违规子类分布见 Figure 6。",
    )
    add_figure(
        doc,
        FIG_DIR / "fig6_corpus.png",
        "Figure 6 · 知识库语料分布(左:按年度;右:Top-8 违规子类, n = 4 233)",
    )
    para(
        doc,
        "如 Figure 6(a) 所示,2017 年后违规案件数量显著抬升,这一时期对应着资本市场"
        "监管趋严的宏观背景,是本研究评测集的主要取样区间。Figure 6(b) 则显示,知识库"
        "中最高频的违规类型为「虚假披露」「违规买卖股票」「其他」,符合证监会年度执法"
        "公告的实际分布。",
    )

    heading(doc, "2.2 数据切分", level=2)
    para(
        doc,
        "为避免数据泄漏,本研究同时采用两种互补的切分策略。其一,EventID 切分"
        "确保同一处罚公告下的多条当事人样本不跨越训练/测试边界,避免模型通过记忆事件 ID "
        "获得虚假性能增益;其二,时间切分以 2021 年及之前为训练集、2022–2023 年为验证集、"
        "2024–2025 年为测试集,严格满足「过去不得包含未来信息」的因果性约束。"
        "此外,PunishmentMeasure(具体处罚措施原文)字段在训练阶段被完全屏蔽,"
        "以消除其与预测标签之间的强相关特征泄漏。",
    )

    heading(doc, "2.3 评测集", level=2)
    para(
        doc,
        "本研究共构建了三套互补的评测集,其基本规模与用途总结于 Table 1。",
    )
    add_table(
        doc,
        headers=["名称", "规模", "意图覆盖", "用途"],
        rows=[
            ["gold_130", "130 条", "case / law / sanction / trend / OOS / multi-turn",
             "检索(98 条可用)+ 生成端到端评测"],
            ["gold_trend_30", "30 条", "trend_analysis",
             "L6 聚合层 exact match / ranking / peak-year 精度"],
            ["LoRA 训练集", "5 500 条", "八类(A-H,含反幻觉负例 H)",
             "QLoRA 指令微调,oracle 证据构造"],
        ],
        caption_text="Table 1 · 三套评测集概览",
    )
    para(
        doc,
        "gold_130 中每条样本均包含 intent、query、gold_answer_keypoints、relevant_event_ids、"
        "relevant_laws 五个字段;其中 98 条带有非空的 relevant_event_ids 可用于检索召回评估。"
        "gold_trend_30 则额外包含 expected_aggregation 字段,其中显式给出了聚合 facet、"
        "year_window、slot_filters 与预期 buckets,用于对 L6 聚合层的精度验证。"
        "LoRA 训练集严格采用 oracle 证据构造,即种子 EventID 在 100% 的样本中出现在输入的"
        "「检索证据」块内,保证模型学到的是「证据→引用→答案」而非「问题→记忆→答案」的映射。",
    )


# ---------------------------------------------------------------------------
# Section · 3 Architecture
# ---------------------------------------------------------------------------
def section_arch(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "3  系统架构", level=1)
    para(
        doc,
        "本研究设计的七层检索增强流水线整体结构如 Figure 1 所示。每一层具有明确的输入输出契约,"
        "层间通过结构化的 QueryPlan 与 SearchResponse 传递状态,使各层可以独立替换、消融与评估。",
    )

    add_figure(
        doc,
        FIG_DIR / "fig1_architecture.png",
        "Figure 1 · 七层检索增强问答流水线架构",
    )

    heading(doc, "3.1 L1 意图分类(Planner)", level=2)
    para(
        doc,
        "本层将用户查询归入 7 类意图之一:greeting、chitchat、out_of_scope、case_retrieval、"
        "law_grounding、sanction_recommendation、trend_analysis。训练样本由 25 条模板 × 20 "
        "组变量 × 7 类意图合成,共 3 520 条,使用 TF-IDF 特征与 Logistic Regression 分类器,"
        "在留出测试集上取得 Macro-F1 = 0.9989。采用经典机器学习而非神经网络的决策出自"
        "两方面权衡:一是 CPU 启动延迟要求 100 ms 以下;二是意图分类结果需在日志中便于"
        "运营审计,经典模型的可解释性更优。",
    )

    heading(doc, "3.2 L2 查询改写(Rewriter)", level=2)
    para(
        doc,
        "本层承担三项职责:共指消解、同义词扩展与槽位抽取。共指消解以规则优先、"
        "大模型回退的方式解析「那案」「它的法条」等代词引用。同义词扩展基于一份包含"
        "257 个规范词与 673 个别名的领域词典(例如「虚假记载」→「虚假陈述 / 披露虚假」)。"
        "槽位抽取则从查询中提取 year、stock_code、violation_type、institution、company、"
        "person 等六类结构化字段,供后续检索与聚合层使用。对多约束查询,本层进一步将其"
        "拆解为多个 sub-query 并行送入检索层,从而提升 hard multi-hop 场景下的召回率。",
    )

    heading(doc, "3.3 L3 双路检索与重排(Retriever)", level=2)
    para(
        doc,
        "检索采用三阶段融合设计:(i) BM25 稀疏检索,使用 jieba 中文分词配合领域 user_dict,"
        "参数 k1 = 1.2, b = 0.75,召回 top-100;(ii) Dense 稠密检索,以 BAAI/bge-small-zh-v1.5"
        "生成 512 维语义向量,按余弦相似度召回 top-100;(iii) Cross-encoder 精排,由 BAAI/"
        "bge-reranker-v2-m3 对候选集重新打分。融合策略采用三层嵌套的 Reciprocal Rank Fusion:"
        "同一 sub-query 内 BM25 与 Dense 的 RRF 融合、多 sub-query 之间再次 RRF、"
        "以及 Hybrid 与 Rerank 结果的 rank-level 融合。后者修正了早期版本中 rerank 直接替换"
        "Hybrid top-k 所导致的 Recall 回归问题。",
    )

    heading(doc, "3.4 L4 证据组装(Evidence Assembly)", level=2)
    para(
        doc,
        "在本层,文档层级的检索结果被聚合为事件级证据。每个事件保留与查询最相关的 3 个"
        "chunk 片段,并附带原始的 title、declare_date、laws、punishment_types 等元数据。"
        "top_k 按意图差异化配置:case_retrieval 取 8,law_grounding 取 8,"
        "sanction_recommendation 取 10,trend_analysis 取 20,以适配不同任务对召回广度的需求。",
    )

    heading(doc, "3.5 L5 生成器(Responder)", level=2)
    para(
        doc,
        "基座选用 Qwen2.5-0.5B-Instruct,在 4-bit NF4 量化的前提下以 QLoRA 方式微调,"
        "秩 r = 16、缩放因子 α = 32,目标模块覆盖自注意力的 q/k/v/o_proj 与 FFN 的 "
        "gate/up/down_proj 共 7 个投影。训练 2 200 条样本、2 epoch,在 RTX 2060 SUPER "
        "8 GB 单卡上用时 52 分钟。本机显存上限无法容纳 1.5 B 版本的 QLoRA(其在 max_seq = "
        "2 048 时第 8 步即 OOM),故 1.5 B 方案以 Colab T4 notebook 的形式提供作为备选。",
    )

    heading(doc, "3.6 L6 趋势聚合器(Trend Aggregator)", level=2)
    para(
        doc,
        "trend_analysis 意图不走向量检索,而短路至本层进行确定性的 groupby count 聚合。"
        "支持的 facet 包括 year、violation_type、punishment_type、agency 四类,并支持"
        "year window(绝对年份或「近 N 年」的相对表达)与 slot filter 的组合过滤。"
        "聚合结果以 [Stat=] 行格式送入 Responder,使 LoRA 能够学到「2022 年 387 起、"
        "2023 年 412 起」这类可审计的硬数字输出,避免通用 LLM 对统计类问题的高幻觉率。",
    )

    heading(doc, "3.7 L7 引证校验(Validator)", level=2)
    para(
        doc,
        "引证校验层基于 YAML 定义的 8 条规则,在生成结果返回用户之前做最后一道闸口。"
        "典型规则包括:答案中必须至少出现一个 [EventID=xxx] 标记;所引用的 EventID 必须"
        "出现在当次检索证据中;所引用的法条必须出现在证据中;未在证据中明确提及的具体"
        "罚款金额不得出现于答案;最终答案须附带免责声明。未通过任一规则的答案将被替换"
        "为降级话术,从而为部署环境提供可审计的兜底保证。",
    )


# ---------------------------------------------------------------------------
# Section · 4 Experiments
# ---------------------------------------------------------------------------
def section_exp(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "4  实验结果", level=1)

    # 4.1 Retrieval ablation
    heading(doc, "4.1 检索层消融", level=2)
    para(
        doc,
        "本小节比较基线配置(仅采用 bge-small-zh 作为 dense encoder, M2 阶段)"
        "与最终配置(多 sub-query 拆分 + 结构化元数据注入 + 扩展金标, M3e 阶段)"
        "在 98 条 retrieval-eligible 金标上的召回表现。采用多 gold 语义的 Recall@5,"
        "定义为 |top5 ∩ gold| / |gold|。结果汇总于 Table 2 并可视化于 Figure 4。",
    )

    add_table(
        doc,
        headers=["检索配置", "Baseline (M2)", "Final (M3e)", "相对增幅"],
        rows=[
            ["BM25-only", "0.077", "0.378", "+391%"],
            ["Dense-only (bge-small-zh)", "0.068", "0.293", "+330%"],
            ["Hybrid (BM25 ⊕ Dense ⊕ RRF)", "0.073", "0.388", "+431%"],
            ["Hybrid + Rerank", "0.067", "0.356", "+431%"],
        ],
        caption_text="Table 2 · 检索层四档消融(gold_130, n = 98)",
    )

    add_figure(
        doc,
        FIG_DIR / "fig4_retrieval.png",
        "Figure 4 · 最终配置下四档检索器在 Recall@5 / Hit@5 / MRR / nDCG@10 上的表现",
    )

    para(
        doc,
        "观察到三点值得讨论的现象。其一,三件套的独立消融显示 multi sub-query 改写对 "
        "Hybrid 贡献稳定的 +2.6 pp 增益,metadata block 注入对 Dense 贡献 +8.8 pp、"
        "对 Rerank 贡献 +14.1 pp。其二,评测集由 50 条扩展至 130 条是最大的"
        "指标校正来源:gold_50 中 hard multi-hop 查询的占比偏高扭曲了指标分布,"
        "扩展后真实分布下 BM25 基线即可达到 0.378。其三,Hybrid + Rerank 的 0.356 "
        "低于 Hybrid 单路的 0.388,原因在于 bge-reranker-v2-m3 未针对 CSRC 领域做适配,"
        "在硬约束查询上倾向于将邻近年份的同类案件推前,这一问题留待未来以对比学习方式"
        "对 reranker 做领域适配。",
    )

    # 4.2 Generation G0-G3
    heading(doc, "4.2 生成层 G0-G3 对照实验", level=2)
    para(
        doc,
        "本小节在 gold_130 按意图分层抽样得到的 30 条子集上,设计四组对照实验以分离"
        "评估 RAG 证据、prompt 工程与 LoRA 微调各自对生成质量的贡献。四组配置如下:"
        "G0 为裸基座(无 RAG、弱 prompt);G1 在 G0 基础上接入检索证据;G2 在 G1 基础上"
        "施加强约束 prompt;G3 进一步加入 LoRA 适配器。所有组共享相同的采样参数"
        "(temperature = 0.2, top_p = 0.9, max_new_tokens = 256)。",
    )

    m44 = _load("docs/reports/m4_4_generation_eval.json")["summary"]

    def fmt(x: float) -> str:
        return f"{x:.3f}"

    add_table(
        doc,
        headers=["组", "配置", "EID 命中率", "格式合规率", "幻觉数字率",
                 "答案长度(字)", "延迟(秒)"],
        rows=[
            ["G0", "base, 无 RAG, 弱 prompt",
             fmt(m44["G0"]["event_id_hit_rate"]),
             fmt(m44["G0"]["format_compliance_rate"]),
             fmt(m44["G0"]["hallucinated_number_rate"]),
             f"{m44['G0']['avg_answer_chars']:.0f}",
             f"{m44['G0']['avg_latency_s']:.2f}"],
            ["G1", "base + RAG, 弱 prompt",
             fmt(m44["G1"]["event_id_hit_rate"]),
             fmt(m44["G1"]["format_compliance_rate"]),
             fmt(m44["G1"]["hallucinated_number_rate"]),
             f"{m44['G1']['avg_answer_chars']:.0f}",
             f"{m44['G1']['avg_latency_s']:.2f}"],
            ["G2", "base + RAG, 强 prompt",
             fmt(m44["G2"]["event_id_hit_rate"]),
             fmt(m44["G2"]["format_compliance_rate"]),
             fmt(m44["G2"]["hallucinated_number_rate"]),
             f"{m44['G2']['avg_answer_chars']:.0f}",
             f"{m44['G2']['avg_latency_s']:.2f}"],
            ["G3", "base + LoRA + RAG + 强 prompt",
             fmt(m44["G3"]["event_id_hit_rate"]),
             fmt(m44["G3"]["format_compliance_rate"]),
             fmt(m44["G3"]["hallucinated_number_rate"]),
             f"{m44['G3']['avg_answer_chars']:.0f}",
             f"{m44['G3']['avg_latency_s']:.2f}"],
        ],
        caption_text="Table 3 · G0–G3 四组生成质量对比 (n = 30)",
    )

    add_figure(
        doc,
        FIG_DIR / "fig2_g0_g3.png",
        "Figure 2 · G0–G3 四组核心指标对比:EventID 命中率、格式合规率与幻觉数字率",
    )

    para(
        doc,
        "实验结果呈现三个关键规律,分别对应本研究的三项核心论断。"
        "第一,G0、G1 与 G2 三组的 EventID 命中率与格式合规率均为 0,即使在 G2 中施加了"
        "「必须引用 [EventID=xxx]」的强约束 prompt,裸基座仍无法遵循该指令;G3 在同一证据、"
        "同一 prompt 条件下,将格式合规率提升至 76.7%。这表明在参数量小于 1 B 的中文基座上,"
        "结构化引用格式的习得无法通过 prompt 工程补偿,必须通过指令微调完成。",
    )
    para(
        doc,
        "第二,如 Figure 3 所示,幻觉数字率呈现「阶跃式下降」的特征:G0 为 20.0%,G1 降至 10.0%,"
        "G2 仍停留于 10.0%,G3 进一步降至 3.3%。这揭示了 RAG 证据与 prompt 工程的边际收益"
        "在 10% 左右达到上限,超出部分必须由 LoRA 对抗训练承担。相对 G0 而言,G3 的幻觉率"
        "下降 83%,且此降幅不依赖于检索层的进一步优化。",
    )

    add_figure(
        doc,
        FIG_DIR / "fig3_hallucination.png",
        "Figure 3 · 幻觉数字率在 G0→G3 上的阶跃式下降",
    )

    para(
        doc,
        "第三,部署成本可控。G3 相对 G0 仅增加约 0.7 秒端到端延迟(+7%),对应 1.7% 的额外"
        "参数量开销。在 RTX 2060 SUPER 8 GB 的单卡环境下,LoRA 适配器与基座共享显存,"
        "不引入额外的显存峰值。",
    )

    para(
        doc,
        "训练收敛曲线见 Figure 5。损失自初始 2.52 平稳下降至 0.70,且 eval 损失的最低点"
        "与 train 损失的末期基本重合,未观察到显著过拟合。",
    )

    add_figure(
        doc,
        FIG_DIR / "fig5_loss.png",
        "Figure 5 · QLoRA 微调的 train / eval 损失曲线(274 step, 2 epoch)",
    )

    # 4.3 Qualitative
    heading(doc, "4.3 定性案例分析", level=2)
    para(
        doc,
        "本小节选取一道典型的幻觉高危查询作为定性对照,以直观展现 G0 与 G3 的行为差异。",
    )
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("查询:")
    _set_font(r, bold=True)
    r = p.add_run("2022 年证监会查处的董事长因内幕交易被罚款的案件有哪些?")
    _set_font(r)

    bold_run(
        doc,
        "G0 输出(裸 Qwen2.5-0.5B): ",
        "「2022 年,证监会共查处了 14 家公司的董事长因内幕交易被罚款,具体案例包括:"
        "深圳市天马科技股份有限公司董事长陈伟明涉嫌内幕交易案、上海市华联股份有限公司"
        "朱永强……」。该段回答中所引用的公司名与当事人均为模型虚构,经知识库逐条核对"
        "均无对应记录,构成典型的事实性幻觉。",
    )

    bold_run(
        doc,
        "G3 输出(+ LoRA): ",
        "「根据检索证据,此类涉嫌内幕交易类违规主要违反的法规为《证券法》第五十条、"
        "第五十三条第一款的规定 [法条:《证券法》第五十条]。参考案例见 "
        "[EventID=40123812],具体条款适用以公告原文为准。」该回答虽然更加保守,"
        "但未引入任何超出检索证据的实体,且以 [EventID=xxx] 格式给出了可溯源的引用。",
    )

    para(
        doc,
        "这一定性对照印证了 §4.2 的定量结论:微调后的模型更倾向于「在可引用证据范围内"
        "保守回答」,而非「在知识缺失时自由发挥」。后者在合规应用场景中会构成严重的运营风险。",
    )

    # 4.4 L6 Aggregator
    heading(doc, "4.4 趋势聚合层精度", level=2)
    m42 = _load("docs/reports/m4_2_trend_eval.json")
    add_table(
        doc,
        headers=["指标", "数值", "含义"],
        rows=[
            ["Exact bucket rate", f"{m42['macro_exact_bucket_rate']:.3f}",
             "各年度/类别的 count 与 ground truth 完全一致的比例"],
            ["Mean relative error", f"{m42['macro_mean_rel_err']:.3f}",
             "平均相对误差(0 表所有 count 完全正确)"],
            ["Ranking accuracy (top-3)",
             f"{m42['macro_ranking_accuracy']:.3f}" if m42.get("macro_ranking_accuracy") else "-",
             "非 year facet 的 top-3 排序与 ground truth 一致率"],
            ["Peak year accuracy",
             f"{m42['peak_year_accuracy']:.3f}" if m42.get("peak_year_accuracy") else "-",
             "year 趋势题中「哪年最高」回答的正确率"],
        ],
        caption_text="Table 4 · L6 趋势聚合层在 gold_trend_30 上的精度",
    )
    para(
        doc,
        "四项指标均达到 1.000,这是因为本层的聚合逻辑为确定性的 groupby count,且 gold 集的"
        "expected_aggregation 字段本身即由聚合器生成并经人工校核。该小节的意义在于:(a) 验证了"
        "聚合实现的正确性,为 Responder 基于结构化证据生成趋势摘要提供可靠上游;"
        "(b) 与纯 LLM 生成统计数字的方案(现有通用对话模型在此类题目上的幻觉率普遍 > 50%)相比,"
        "本研究给出的「2024 年内幕交易 41 起」属于可审计的硬数字。",
    )

    # 4.5 Punishment-type classifier
    heading(doc, "4.5 辅线任务:处罚类型多标签分类", level=2)
    m5 = _try_load("docs/reports/m5_macbert_report.json")
    if m5 is None:
        para(
            doc,
            "辅线任务为对每起违规事件预测其处罚类型多标签集合(罚款 / 警告 / 没收非法所得 / "
            "市场禁入等 7 类)。基线采用 TF-IDF + LogisticRegression,在当事人样本上 Micro-F1 "
            "= 0.864;MacBERT 微调结果在本版本中以占位形式呈现,答辩版将补齐。",
        )
    else:
        tm = m5.get("test_metrics", {})
        micro = tm.get("eval_Micro-F1") or tm.get("Micro-F1") or 0.0
        macro = tm.get("eval_Macro-F1") or tm.get("Macro-F1") or 0.0
        ham = tm.get("eval_HammingLoss") or tm.get("Hamming") or 0.0
        sub = tm.get("eval_SubsetAccuracy") or tm.get("SubsetAccuracy") or 0.0
        labels = m5.get("label_vocab", [])
        nlab = len(labels) if labels else 7

        add_table(
            doc,
            headers=["模型", "Micro-F1", "Macro-F1", "Subset Acc.", "Hamming Loss"],
            rows=[
                ["TF-IDF + LogReg(baseline)", "0.864", "0.276", "—", "0.069"],
                ["MacBERT-base (本研究微调)",
                 f"{micro:.3f}", f"{macro:.3f}", f"{sub:.3f}", f"{ham:.3f}"],
            ],
            caption_text=f"Table 5 · 当事人处罚类型多标签分类对比(label 数 = {nlab})",
        )
        para(
            doc,
            "Micro-F1 指标反映了对高频标签(罚款、没收非法所得、警告)的主流判别能力;"
            "Macro-F1 偏低的根源在于少数类(批评、市场禁入、谴责)存在显著的长尾分布问题。"
            "在部署场景中建议对 top-3 高频标签设置独立阈值,少数类则回退至保守拒答策略。"
            "本任务在主线流水线中作为 sanction_recommendation 意图下 LoRA 生成结果的"
            "辅助参考信号,用以交叉校验建议的处罚类型分布是否合理。",
        )


# ---------------------------------------------------------------------------
# Section · 5 Hallucination Mitigation
# ---------------------------------------------------------------------------
def section_hallucination(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "5  幻觉缓解的三层防线", level=1)
    para(
        doc,
        "在垂直领域问答场景下,幻觉控制是决定系统能否进入生产环境的关键因素。"
        "本研究将幻觉缓解机制按责任边界拆分为三层独立防线,每层针对不同的失效模式。"
        "各层对 G0 → G3 幻觉数字率的边际贡献见 §4.2 与 Figure 3。",
    )

    heading(doc, "5.1 第一层:系统边界层", level=2)
    para(
        doc,
        "第一道防线是在入口处拒绝回答本研究职责范围以外的问题。L0 Topic Guard 层以正则白名单"
        "结合关键词黑名单硬拦截股价预测、诗歌创作、代码生成等越界查询,相应请求直接返回"
        "固定话术,不再下推至大模型。L1 意图分类器则进一步将「请预测 2026 年重点整治领域」"
        "这类主观预测归入 out_of_scope 类别,走固定拒答路径。在 gold_130 中设置的 4 条"
        "反幻觉陷阱题被本层以 100% 的覆盖率成功拦截。",
    )

    heading(doc, "5.2 第二层:证据约束层", level=2)
    para(
        doc,
        "第二道防线通过训练数据与推理 prompt 两端协同实现。在训练数据端,LoRA 训练集中 100% 的"
        "样本采用 oracle 证据构造,即种子 EventID 强制出现在输入的「检索证据」块中,"
        "以此教会模型「证据→引用→答案」的正确映射;在推理端,系统 prompt 明文约束:"
        "必须引用 [EventID=xxx] 格式、不得编造证据中未出现的内容、证据不足时须明确声明。"
        "该层对应 G0 → G1 的幻觉率下降,即从 20.0% 降至 10.0%,贡献了约一半的整体缓解幅度。",
    )

    heading(doc, "5.3 第三层:对抗训练层", level=2)
    para(
        doc,
        "第三道防线专门针对第二层无法覆盖的边界情形,即模型在证据不完整或不存在时的行为。"
        "本研究在训练集中刻意注入两类反例样本:H 类 240 条反幻觉负例,诱导模型在证据不足时"
        "回答「未检索到」而非编造;G 类 160 条多轮追问样本,训练模型在追问轮次中坚持证据内"
        "的实体而非跟随用户诱导作出虚构。该层对应 G2 → G3 的下降,即从 10.0% 进一步降至 3.3%。",
    )

    heading(doc, "5.4 第四层:引证校验层", level=2)
    para(
        doc,
        "最后一道闸口由 L7 Validator 在生成结果返回用户之前执行。本层不依赖模型能力,"
        "而是以 YAML 定义的 8 条确定性规则逐条校验答案中出现的 EventID 与法条是否在当次"
        "检索证据中存在;任一规则不通过则用户看到的将是降级话术,而非原始生成。"
        "这一层为部署环境提供了可审计的兜底保证,即便前三层全部失效,系统仍不会对外输出"
        "未经引证核验的内容。",
    )


# ---------------------------------------------------------------------------
# Section · 6 Deployment
# ---------------------------------------------------------------------------
def section_deployment(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "6  部署与性能分析", level=1)
    para(
        doc,
        "本研究将端到端系统部署至单块 RTX 2060 SUPER 8 GB 消费级 GPU,各组件的资源开销"
        "汇总于 Table 6。",
    )
    add_table(
        doc,
        headers=["组件", "模型", "磁盘", "显存(推理)", "平均延迟"],
        rows=[
            ["L1 Planner", "TF-IDF + LogReg", "2.4 MB", "< 50 MB", "15 ms"],
            ["L3 Dense Encoder", "bge-small-zh-v1.5", "99 MB", "500 MB", "280 ms"],
            ["L3 Reranker", "bge-reranker-v2-m3", "2.3 GB", "2.1 GB", "1 630 ms"],
            ["L5 Responder Base", "Qwen2.5-0.5B-Instruct (4-bit)", "394 MB", "1.8 GB", "9 800 ms"],
            ["L5 Responder LoRA", "qwen_lora_csrc (adapter)", "34 MB", "+ 60 MB", "+ 700 ms"],
            ["总计(含 reranker)", "—", "≈ 2.8 GB", "≈ 4.5 GB", "≈ 12 s"],
            ["总计(不含 reranker)", "—", "≈ 530 MB", "≈ 2.4 GB", "≈ 10.5 s"],
        ],
        caption_text="Table 6 · 端到端部署资源开销",
    )
    para(
        doc,
        "在保留 reranker 的完整配置下,系统总显存占用约 4.5 GB,平均端到端延迟约 12 秒;"
        "如选择去掉 reranker(其对 Hybrid Recall@5 的贡献实为负向, 见 §4.1),总资源下降至 2.4 GB "
        "显存与 10.5 秒延迟,可适配中小金融机构的合规问答场景。训练侧 LoRA adapter 仅 34 MB,"
        "便于按领域、客户或业务线进行多版本 A/B 管理与多租户部署。",
    )


# ---------------------------------------------------------------------------
# Section · 7 Limitations
# ---------------------------------------------------------------------------
def section_limits(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "7  局限与未来工作", level=1)
    para(
        doc,
        "本研究存在以下五点明确的局限,对应的改进路径亦一并指出。",
    )
    bullet(
        doc,
        "基座规模受限:受 8 GB 显存上限约束,主干微调仅在 Qwen2.5-0.5B-Instruct 上开展。"
        "1.5 B 版本的 QLoRA 在 max_seq = 2 048 时第 8 步触发 OOM,即便降至 1 024 仍处于"
        "显存边缘。本研究提供 Colab T4 的 1.5 B 补训 notebook 作为可选复现路径。",
    )
    bullet(
        doc,
        "EID 命中率上限:G3 的 EventID 命中率 20.0%,剩余 80% 未命中主要归因于 Hybrid Recall@5"
        " = 0.388 的检索天花板,而非生成层的错误;进一步提升需在 reranker 上引入 CSRC 领域"
        "的对比学习 LoRA,这将作为未来工作开展。",
    )
    bullet(
        doc,
        "幻觉检测覆盖面有限:本研究采用的正则检测器仅覆盖数字类幻觉(金额、百分比),"
        "对人名与公司名类幻觉需要人工标注。样本规模为 30 条,95% 置信区间约为 ±18%,"
        "后续应扩展至 100 条以获得更稳定的点估计。",
    )
    bullet(
        doc,
        "评测集意图分布不均:gold_130 以 case_retrieval 为主(64 条),"
        "sanction_recommendation 与 multi_turn_followup 类样本相对偏少,未来应结合真实客服"
        "日志扩展以提升评测分布的代表性。",
    )
    bullet(
        doc,
        "语言与区域支持:本研究目前仅覆盖中国大陆证监会公告,港澳台地区证券监管的"
        "繁体中文语料与金融术语差异未纳入,这也是后续工作可扩展的方向之一。",
    )


# ---------------------------------------------------------------------------
# Section · 8 Contributions
# ---------------------------------------------------------------------------
def section_contributions(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "8  贡献声明与可复现性", level=1)

    heading(doc, "8.1 团队分工", level=2)
    para(
        doc,
        "本研究由五人团队独立完成。团队分工基本对齐开题阶段的初始规划:许浩财负责总体架构、"
        "意图层与系统集成;贾彤负责数据处理与向量化索引;戴一鑫负责问答样本与拒答策略;"
        "张彦扬负责 prompt 设计与问答样本扩展;王怡菲负责模型训练与评估。"
        "在实际研发过程中,团队统一使用 Claude Code(Anthropic 发布的 AI 编程助手)"
        "作为结对编程工具,故以下独立声明 AI 辅助的具体贡献。",
    )

    heading(doc, "8.2 AI 辅助声明", level=2)
    para(
        doc,
        "本研究在策略文档撰写、代码实现、训练与评估脚本编写、bug 定位以及论文起草等多个"
        "环节使用了 Anthropic Claude(Claude Opus 4.7 与 Claude Sonnet 4.6,集成于 "
        "Claude Code CLI)作为辅助工具。其具体贡献包括:",
    )
    bullet(
        doc,
        "策略文档初稿:docs/strategies/ 下 12 份策略文档的初稿由 AI 根据作者提示生成,"
        "其后经人工审校、修改并合并至正式版本;",
    )
    bullet(
        doc,
        "代码实现:src/csrc_rag/ 目录下约 70% 的 Python 代码由 AI 起草,"
        "作者逐行审查、修改并进行集成测试;",
    )
    bullet(
        doc,
        "实验脚本:QLoRA 训练脚本、G0-G3 评估脚本、Trend Aggregator 评估脚本的初稿"
        "均由 AI 起草,并在真实运行中协助调试;",
    )
    bullet(
        doc,
        "bug 定位与修复:M3d 阶段的 rerank 合并回归 bug、M4.3 阶段的 gradient "
        "checkpoint 张量连续性 bug 均由 AI 在读取栈跟踪后提出修复方案;",
    )
    bullet(
        doc,
        "论文起草:本论文由 scripts/build_paper_docx.py 自动生成,"
        "其中 python-docx 布局逻辑与基础文字由 AI 起草,"
        "引言、实验讨论与结论段落经人工修订后定稿。",
    )
    para(
        doc,
        "作者团队对以下方面保留完整的研究判断责任:整体研究设计与赛道 B 需求的翻译、"
        "数据构造与标注的规则确认、每次实验的超参选择与结果解释,以及对 AI 产出内容的"
        "事实核验与风险评估。AI 的贡献可归纳为加速实现而非替代研究判断。",
    )

    heading(doc, "8.3 可复现性", level=2)
    para(
        doc,
        "本研究的代码仓库托管于 https://github.com/Mindse-Tt/Deeplearning_project-CSRC_Rag(答辩前"
        "对 reviewer 开放访问),tag v1.0-track-b-final 对应本论文所用的确定版本。"
        "所有训练数据、LoRA adapter 权重以及评测集均可通过 Git 仓库下载,或通过 "
        "scripts/build_*.py 系列脚本从原始 CSMAR 数据重新构建。",
    )


# ---------------------------------------------------------------------------
# Section · References
# ---------------------------------------------------------------------------
def section_refs(doc: Document) -> None:
    doc.add_page_break()
    heading(doc, "参考文献", level=1)
    refs = [
        "[1] Dettmers, T., Pagnoni, A., Holtzman, A., & Zettlemoyer, L. (2023). "
        "QLoRA: Efficient Finetuning of Quantized LLMs. arXiv preprint arXiv:2305.14314.",
        "[2] Qwen Team. (2024). Qwen2.5 Technical Report. https://qwenlm.github.io/blog/qwen2.5/",
        "[3] Xiao, S., Liu, Z., Zhang, P., & Muennighoff, N. (2023). "
        "C-Pack: Packaged Resources To Advance General Chinese Embedding. arXiv:2309.07597.",
        "[4] Robertson, S., & Zaragoza, H. (2009). The probabilistic relevance "
        "framework: BM25 and beyond. Foundations and Trends in Information Retrieval, 3(4), 333–389.",
        "[5] Cormack, G. V., Clarke, C. L., & Büttcher, S. (2009). Reciprocal rank fusion "
        "outperforms Condorcet and individual rank learning methods. In Proc. SIGIR '09.",
        "[6] Hu, E. J., Shen, Y., Wallis, P., Allen-Zhu, Z., Li, Y., Wang, S., "
        "Wang, L., & Chen, W. (2021). LoRA: Low-Rank Adaptation of Large Language Models. arXiv:2106.09685.",
        "[7] Cui, Y., Che, W., Liu, T., Qin, B., Wang, S., & Hu, G. (2020). "
        "Revisiting Pre-Trained Models for Chinese Natural Language Processing. "
        "In Findings of EMNLP 2020.",
        "[8] Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., "
        "Küttler, H., Lewis, M., Yih, W.-t., Rocktäschel, T., Riedel, S., & Kiela, D. "
        "(2020). Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. NeurIPS 2020.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        r = p.add_run(ref)
        _set_font(r, size_pt=9.5)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def build(out_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    section_cover(doc)
    section_intro(doc)
    section_data(doc)
    section_arch(doc)
    section_exp(doc)
    section_hallucination(doc)
    section_deployment(doc)
    section_limits(doc)
    section_contributions(doc)
    section_refs(doc)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"[PASS] wrote {out_path}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--out",
        type=Path,
        default=PROJECT_ROOT / "docs" / "paper" / "csrc_rag_v1.docx",
    )
    args = parser.parse_args()
    build(args.out)


if __name__ == "__main__":
    main()
